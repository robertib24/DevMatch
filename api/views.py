from rest_framework import viewsets, status, filters, pagination
from rest_framework.decorators import action, api_view
from rest_framework.response import Response
from django.core.files.uploadedfile import UploadedFile
from .models import CV, JobDescription, MatchResult
from .serializers import CVSerializer, JobDescriptionSerializer, MatchResultSerializer
from .gemini_utils import GeminiAI
from django.db import transaction
from django.utils import timezone
import os
import re
import docx
import logging
from datetime import datetime
from django.core.cache import cache

logger = logging.getLogger(__name__)

class StandardResultsSetPagination(pagination.PageNumberPagination):
    page_size = 10
    page_size_query_param = 'page_size'
    max_page_size = 100

class CVViewSet(viewsets.ModelViewSet):
    queryset = CV.objects.all()
    serializer_class = CVSerializer
    pagination_class = StandardResultsSetPagination
    filter_backends = [filters.SearchFilter, filters.OrderingFilter]
    search_fields = ['name', 'content']
    ordering_fields = ['name', 'processed_at']
    ordering = ['-processed_at']

    def create(self, request, *args, **kwargs):
        print("--- Entering CVViewSet CREATE method ---")
        file: UploadedFile | None = request.FILES.get('file')
        if not file:
            print("!!! CV Upload Error: No file provided.")
            return Response({'detail': 'No file provided.'}, status=status.HTTP_400_BAD_REQUEST)

        name = request.data.get('name', file.name)
        print(f"Processing CV file: {name}")

        if not file.name.lower().endswith('.docx'):
            print(f"!!! CV Upload Error: Invalid file type for {name}.")
            return Response({'detail': 'Invalid file type. Only .docx allowed.'}, status=status.HTTP_400_BAD_REQUEST)

        try:
            file.seek(0)
            print(f"Attempting to read docx content for {name}...")
            document = docx.Document(file)
            paras = [p.text for p in document.paragraphs if p.text.strip()]
            content = "\n".join(paras)
            print(f"Extracted content length for {name}: {len(content)}")
            if not content:
                print(f"!!! CV Upload Error: Could not extract text from {name}.")
                return Response({'detail': 'Could not extract text content from the file.'}, status=status.HTTP_400_BAD_REQUEST)

        # Use transaction to ensure database consistency
            with transaction.atomic():
                file.seek(0)
                print(f"Attempting to create CV object for {name}...")
                cv = CV.objects.create(name=name, file=file, content=content)
                print(f"CV object created with ID: {cv.id}")

            serializer = self.get_serializer(cv)
            print(f"Successfully processed CV {name}.")
            return Response(serializer.data, status=status.HTTP_201_CREATED)

        except Exception as e:
            print(f"!!! CV Upload Error: Exception during processing {name}: {e}")
            # import traceback
            # traceback.print_exc()
            return Response({'detail': f'Error processing CV file: {str(e)}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


    @action(detail=True, methods=['get'])
    def find_best_job(self, request, pk=None):
        try:
            cv = self.get_object()
        except CV.DoesNotExist:
             return Response({'message': 'CV not found.'}, status=status.HTTP_404_NOT_FOUND)
        
        # Get all jobs, with optional filtering
        jobs_queryset = JobDescription.objects.all()
        industry_filter = request.query_params.get('industry')
        if industry_filter:
            jobs_queryset = jobs_queryset.filter(industry__icontains=industry_filter)
            
        jobs = jobs_queryset
        
        if not jobs: 
            return Response({'message': 'No jobs found'}, status=status.HTTP_404_NOT_FOUND)
            
        try: 
            gemini = GeminiAI()
        except Exception as e:
             print(f"Error initializing GeminiAI: {e}")
             return Response({'message': 'AI service configuration error.'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
             
        # Check if we already have matches in the database to avoid reprocessing
        existing_matches = MatchResult.objects.filter(cv=cv).order_by('-total_score')
        
        if existing_matches.exists():
            # Use existing top match
            top_match = existing_matches.first()
            
            # Return the top match data
            return Response({
                'job': JobDescriptionSerializer(top_match.job).data,
                'total_score': top_match.total_score,
                'industry_score': top_match.industry_score,
                'tech_skills_score': top_match.tech_skills_score,
                'description_match_score': top_match.description_match_score,
                'explanation': getattr(top_match, 'explanation', "Match retrieved from database")
            }, status=status.HTTP_200_OK)
        
        # If no existing matches, process all jobs
        # This is inefficient for many jobs - in a real-world scenario, 
        # you might want to limit the number of jobs to process
        results = []
        # Prepare batch processing
        batch_matches = []
        
        for job in jobs:
            batch_matches.append({
                'cv_id': cv.id,
                'job_id': job.id,
                'cv_text': cv.content,
                'job_description': job.content,
                'job_industry': job.industry,
                'technical_skills': job.technical_skills or {}
            })
        
        # Process in batch (more efficient)
        if batch_matches:
            batch_results = gemini.process_matches_parallel(batch_matches)
            
            # Process all results and save to database
            for match_data in batch_results:
                job_id = match_data['job_id']
                match_result = match_data['result']
                job = JobDescription.objects.get(id=job_id)
                
                # Calculate total score
                total_score = (
                    match_result.get('industry_score', 0) * 0.1 +
                    match_result.get('tech_skills_score', 0) * 0.3 +
                    match_result.get('description_match_score', 0) * 0.6
                )
                
                # Create/update match in database
                match, _ = MatchResult.objects.update_or_create(
                    cv=cv, job=job,
                    defaults={
                        'industry_score': match_result.get('industry_score', 0),
                        'tech_skills_score': match_result.get('tech_skills_score', 0),
                        'description_match_score': match_result.get('description_match_score', 0),
                        'total_score': total_score, 
                        'matched_at': timezone.now(),
                        'explanation': match_result.get('explanation', '')
                    }
                )
                
                # Add to results list
                results.append({
                    'job': JobDescriptionSerializer(job).data,
                    'total_score': total_score,
                    'industry_score': match_result.get('industry_score', 0),
                    'tech_skills_score': match_result.get('tech_skills_score', 0),
                    'description_match_score': match_result.get('description_match_score', 0),
                    'explanation': match_result.get('explanation', '')
                })
        
        # Sort results by total score
        results.sort(key=lambda x: x['total_score'], reverse=True)
        return Response(results[0] if results else {'message': 'No suitable match found.'}, 
                       status=status.HTTP_200_OK if results else status.HTTP_404_NOT_FOUND)

class JobDescriptionViewSet(viewsets.ModelViewSet):
    queryset = JobDescription.objects.all()
    serializer_class = JobDescriptionSerializer
    pagination_class = StandardResultsSetPagination
    filter_backends = [filters.SearchFilter, filters.OrderingFilter]
    search_fields = ['title', 'content', 'industry']
    ordering_fields = ['title', 'industry', 'created_at']
    ordering = ['-created_at']

    def create(self, request, *args, **kwargs):
        print("--- Entering JobDescriptionViewSet CREATE method ---")
        file: UploadedFile | None = request.FILES.get('file')
        if not file:
            print("!!! Job Upload Error: No file provided.")
            return Response({'detail': 'No file provided.'}, status=status.HTTP_400_BAD_REQUEST)

        print(f"Processing Job Description file: {file.name}")

        if not file.name.lower().endswith('.docx'):
            print(f"!!! Job Upload Error: Invalid file type for {file.name}.")
            return Response({'detail': 'Invalid file type. Only .docx allowed.'}, status=status.HTTP_400_BAD_REQUEST)

        try:
            # 1. Extract text
            try:
                file.seek(0)
                print(f"Attempting to read docx content for {file.name}...")
                document = docx.Document(file)
                paras = [p.text for p in document.paragraphs if p.text.strip()]
                full_text = "\n".join(paras)
                print(f"Extracted job content length: {len(full_text)}")
                if not full_text:
                    print(f"!!! Job Upload Error: Could not extract text from {file.name}.")
                    return Response({'detail': 'Could not extract text content from the file.'}, status=status.HTTP_400_BAD_REQUEST)
            except Exception as e:
                print(f"!!! Job Upload Error: Failed reading docx {file.name}: {e}")
                return Response({'detail': f'Failed to read docx file content: {e}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

            # 2. Parse the extracted text
            print(f"Attempting to parse text for {file.name}...")
            parsed_data = self._parse_job_description_text(full_text)
            parsed_data['content'] = full_text  # Save the full text
            print(f"Parsed Data for {file.name}: Title='{parsed_data.get('title')}', Industry='{parsed_data.get('industry')}', Skills Count={len(parsed_data.get('technical_skills', {}))}")

            # 3. Check if title was parsed
            if not parsed_data.get('title'):
                print(f"!!! Job Upload Error: Could not parse title for {file.name}.")
                return Response({'detail': 'Could not parse job title from the document.'}, status=status.HTTP_400_BAD_REQUEST)

            # 4. Create or update JobDescription object with transaction
            print(f"Attempting to update or create JobDescription for '{parsed_data['title']}'...")
            with transaction.atomic():
                job, created = JobDescription.objects.update_or_create(
                    title=parsed_data['title'],
                    defaults={
                    'content': parsed_data['content'],
                    'industry': parsed_data['industry'],
                    'technical_skills': parsed_data['technical_skills']
                }
            )
            print(f"JobDescription {'created' if created else 'updated'} with ID: {job.id}")

            # 5. Serialize and return response
            serializer = self.get_serializer(job)
            status_code = status.HTTP_201_CREATED if created else status.HTTP_200_OK
            print(f"Successfully processed job description {file.name}.")
            return Response(serializer.data, status=status_code)

        except Exception as e:
            print(f"!!! Job Upload Error: General exception during processing {file.name}: {e}")
            return Response({'detail': f'Error processing job description file: {str(e)}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        
    @action(detail=True, methods=['get'])
    def top_candidates(self, request, pk=None):
        """Get top 5 candidates for a specific job with optimized response time"""
        try: 
            job = self.get_object()
        except JobDescription.DoesNotExist: 
            return Response({'message': 'Job not found.'}, status=status.HTTP_404_NOT_FOUND)
        
        # Get the top 5 matches for this job
        # Use database-level filtering for efficiency
        top_matches = MatchResult.objects.filter(job=job).order_by('-total_score')[:5]
    
        if not top_matches:
            return Response([], status=status.HTTP_200_OK)
        
        # We can optimize by using select_related to reduce database queries
        top_matches = MatchResult.objects.filter(job=job).select_related('cv').order_by('-total_score')[:5]
        
        results = []
        for match in top_matches:
            results.append({
                'cv': CVSerializer(match.cv).data,
                'total_score': match.total_score,
                'industry_score': match.industry_score,
                'tech_skills_score': match.tech_skills_score,
                'description_match_score': match.description_match_score,
                'explanation': getattr(match, 'explanation', '')
            })
    
        return Response(results, status=status.HTTP_200_OK)

    def _parse_job_description_text(self, text):
        data = {
            'title': None,
            'industry': 'Not Specified',
            'technical_skills': {},
        }
        lines = text.splitlines()
        potential_skills = []
        in_skills_section = False

        print("--- Starting text parsing (V2) ---")

        # Extract Title (look for 'Job Title:' and take the next line)
        for i, line in enumerate(lines):
            line_stripped = line.strip()
            if line_stripped.lower().startswith("job title:"):
                print(f"Found 'Job Title:' marker at line {i+1}")
                # Look for the next non-empty line
                if i + 1 < len(lines):
                    next_line_stripped = lines[i+1].strip()
                    if next_line_stripped:
                        data['title'] = next_line_stripped
                        print(f"Found Title on next line: '{data['title']}'")
                        break
        # If we didn't find the title after marker, try a heuristic (first line?)
        if not data['title'] and lines:
             first_line = lines[0].strip()
             # Consider the first line as title if it's not too long and not a generic header
             if len(first_line.split()) < 10 and not first_line.lower().startswith("company overview"):
                   print(f"Using first line as fallback title: '{first_line}'")
                   data['title'] = first_line


        # Extract Industry
        for line in lines:
            line_stripped = line.strip()
            if line_stripped.lower().startswith("industry:"):
                 data['industry'] = line_stripped[len("industry:"):].strip()
                 print(f"Found Industry: {data['industry']}")
                 break

        # Extract Skills
        print("Scanning for skills sections...")
        for line in lines:
            line_stripped = line.strip()
            line_lower = line_stripped.lower()

            # Detect section starts
            if line_lower.startswith("required qualifications") or \
               line_lower.startswith("qualifications:") or \
               line_lower.startswith("preferred skills") or \
               line_lower.startswith("skills:") or \
               line_lower.startswith("technical skills") or \
               line_lower.startswith("requirements"):
                if not in_skills_section:
                    print(f"Entering potential skills section at line: '{line_stripped}'")
                    in_skills_section = True
                continue # Skip the header line

            # Detect section ends / other major sections
            # Check if we ARE in a section BEFORE exiting
            if in_skills_section and (
               line_lower.startswith("benefits:") or \
               line_lower.startswith("key responsibilities:") or \
               line_lower.startswith("responsibilities:") or \
               line_lower.startswith("company overview:") or \
               line_lower.startswith("about us:") or \
               line_lower.startswith("job description:") or \
               (line_stripped.endswith(':') and len(line_stripped.split()) < 5) ):
                 print(f"Exiting potential skills section at line: '{line_stripped}'")
                 in_skills_section = False
                 # Skip this line, it's part of a new section

            # Add skill if we are in a section and the line seems valid
            if in_skills_section and line_stripped:
                skill_match = re.match(r'^[-*•\u2022\s]+(.*)', line_stripped)
                skill = ""
                if skill_match:
                    skill = skill_match.group(1).strip()
                elif len(line_stripped.split()) < 15:
                     # Heuristic for lines without a marker, only if they don't appear to be titles
                     if not line_stripped.endswith(':'):
                          skill = line_stripped

                skill = skill.rstrip('.').strip()
                if skill and len(skill) > 1:
                    if not (skill.endswith('.') or skill.endswith(':')):
                         potential_skills.append(skill)

        # Process skills
        unique_skills = sorted(list(set(potential_skills)))
        data['technical_skills'] = {skill: 100 for skill in unique_skills}
        print(f"Final parsed skills: {data['technical_skills']}")

        print("--- Finished text parsing (V2) ---")
        return data


    @action(detail=True, methods=['get'])
    def find_matching_cvs(self, request, pk=None):
        """Find matching CVs for a job with improved performance and parallel processing."""
        try: 
            job = self.get_object()
        except JobDescription.DoesNotExist: 
            return Response({'message': 'Job not found.'}, status=status.HTTP_404_NOT_FOUND)
        
        # Get parameters for potential filtering
        min_score = request.query_params.get('min_score', 0)
        try:
            min_score = float(min_score)
        except ValueError:
            min_score = 0
            
        # Get CVs with optional filtering
        cv_queryset = CV.objects.all()
        
        # Get name filter if provided
        name_filter = request.query_params.get('name')
        if name_filter:
            cv_queryset = cv_queryset.filter(name__icontains=name_filter)
            
        # Check if we should use cached results
        use_cache = request.query_params.get('use_cache', 'true').lower() == 'true'
        
        # Get max number of CVs to process
        max_cvs = int(request.query_params.get('max_cvs', 10))
        cvs = cv_queryset[:max_cvs]
        
        if not cvs: 
            return Response({'message': 'No CVs found matching criteria'}, status=status.HTTP_404_NOT_FOUND)
            
        try: 
            gemini = GeminiAI()
        except Exception as e:
             print(f"Error initializing GeminiAI: {e}")
             return Response({'message': 'AI service configuration error.'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
             
        # Prepare for batch processing
        batch_matches = []
        for cv in cvs:
            if not cv.content:
                print(f"Skipping CV {cv.id} due to missing content.")
                continue
                
            if not use_cache:
                # If not using cache, we need to check if there's an existing match and delete it
                existing_match = MatchResult.objects.filter(job=job, cv=cv).first()
                if existing_match:
                    existing_match.delete()
                    
            # Add to batch for processing
            batch_matches.append({
                'cv_id': cv.id,
                'job_id': job.id,
                'cv_text': cv.content,
                'job_description': job.content,
                'job_industry': job.industry,
                'technical_skills': job.technical_skills or {}
            })
        
        results = []
        
        # First check if we have existing matches in the database (if using cache)
        if use_cache:
            existing_matches = MatchResult.objects.filter(
                job=job, 
                cv__in=cvs
            ).select_related('cv')
            
            # Add existing matches to results
            for match in existing_matches:
                results.append({
                    'cv': CVSerializer(match.cv).data, 
                    'total_score': match.total_score,
                    'industry_score': match.industry_score,
                    'tech_skills_score': match.tech_skills_score,
                    'description_match_score': match.description_match_score,
                    'explanation': getattr(match, 'explanation', "Previously analyzed match")
                })
            
            # Filter out CVs that already have matches
            processed_cv_ids = [match.cv.id for match in existing_matches]
            batch_matches = [m for m in batch_matches if m['cv_id'] not in processed_cv_ids]
        
        # Process remaining CVs in parallel
        if batch_matches:
            # Use improved parallel processing
            batch_results = gemini.process_matches_parallel(batch_matches)
            
            for match_data in batch_results:
                cv_id = match_data['cv_id']
                match_result = match_data['result']
                
                # Get the CV
                try:
                    cv = CV.objects.get(id=cv_id)
                    
                    # Calculate total score
                    total_score = (
                        match_result.get('industry_score', 0) * 0.1 +
                        match_result.get('tech_skills_score', 0) * 0.3 +
                        match_result.get('description_match_score', 0) * 0.6
                    )
                    
                    # Skip if below minimum score threshold
                    if total_score < min_score:
                        continue
                    
                    # Get the explanation
                    explanation = match_result.get('explanation', '')
                    
                    # Save the match to database with explanation
                    match, _ = MatchResult.objects.update_or_create(
                        cv=cv, job=job,
                        defaults={
                            'industry_score': match_result.get('industry_score', 0),
                            'tech_skills_score': match_result.get('tech_skills_score', 0),
                            'description_match_score': match_result.get('description_match_score', 0),
                            'total_score': total_score, 
                            'matched_at': timezone.now(),
                            'explanation': explanation
                        }
                    )
                    
                    # Add to results
                    results.append({
                        'cv': CVSerializer(match.cv).data,
                        'total_score': total_score,
                        'industry_score': match_result.get('industry_score', 0),
                        'tech_skills_score': match_result.get('tech_skills_score', 0),
                        'description_match_score': match_result.get('description_match_score', 0),
                        'explanation': explanation
                    })
                except CV.DoesNotExist:
                    continue
        
        # Sort results by total score
        results.sort(key=lambda x: x['total_score'], reverse=True)
        
        # Return top 5 results with explanations
        top_5 = results[:5]
        return Response(top_5)

class MatchResultViewSet(viewsets.ReadOnlyModelViewSet):
    """
    ViewSet for viewing match results between CVs and job descriptions.
    """
    # Use select_related to optimize database queries
    queryset = MatchResult.objects.all().select_related('cv', 'job').order_by('-total_score')
    serializer_class = MatchResultSerializer
    pagination_class = StandardResultsSetPagination
    filter_backends = [filters.SearchFilter, filters.OrderingFilter]
    search_fields = ['cv__name', 'job__title']
    ordering_fields = ['total_score', 'industry_score', 'tech_skills_score', 'description_match_score', 'matched_at']
    ordering = ['-total_score']
    
    def get_queryset(self):
        """
        Return filtered match results based on query parameters.
        """
        queryset = MatchResult.objects.all().select_related('cv', 'job')
        
        # Filter by minimum score if provided
        min_score = self.request.query_params.get('min_score')
        if min_score:
            try:
                min_score_float = float(min_score)
                queryset = queryset.filter(total_score__gte=min_score_float)
            except ValueError:
                pass
                
        # Filter by industry if provided
        industry = self.request.query_params.get('industry')
        if industry:
            queryset = queryset.filter(job__industry__icontains=industry)
            
        # Filter by job title if provided
        job_title = self.request.query_params.get('job_title')
        if job_title:
            queryset = queryset.filter(job__title__icontains=job_title)
            
        # Filter by CV name if provided
        cv_name = self.request.query_params.get('cv_name')
        if cv_name:
            queryset = queryset.filter(cv__name__icontains=cv_name)
            
        # Filter by date range if provided
        start_date = self.request.query_params.get('start_date')
        end_date = self.request.query_params.get('end_date')
        if start_date:
            queryset = queryset.filter(matched_at__gte=start_date)
        if end_date:
            queryset = queryset.filter(matched_at__lte=end_date)
        
        # Apply custom ordering with fallback to default
        ordering = self.request.query_params.get('ordering', '-total_score')
        return queryset.order_by(ordering)

@api_view(['GET'])
def get_statistics(request):
    """
    Get system statistics with optimized database queries and caching.
    """
    cache.clear()
    # Try to get statistics from cache
    stats_cache_key = 'system_statistics'
    cached_stats = cache.get(stats_cache_key)
    
    if cached_stats:
        # Add recent activities since they change more frequently
        recent_activities = get_recent_activities()
        cached_stats['recentActivity'] = recent_activities
        return Response(cached_stats)
    
    # If not in cache, calculate statistics using efficient database queries
    stats = {
        'totalCVs': CV.objects.count(),
        'totalJobs': JobDescription.objects.count(),
        'totalMatches': MatchResult.objects.count(),
        'recentActivity': []
    }
    
    # Add more useful statistics
    try:
        # Get recent activities
        stats['recentActivity'] = get_recent_activities()
        
        # Cache the basic statistics for 10 minutes
        # We exclude recentActivity since that changes frequently
        stats_to_cache = {k: v for k, v in stats.items() if k != 'recentActivity'}
        cache.set(stats_cache_key, stats_to_cache, 600)  # 10 minutes
        
    except Exception as e:
        print(f"Error calculating additional statistics: {e}")
    
    return Response(stats)

def get_recent_activities():
    """Helper function to get recent activities for the statistics endpoint."""
    recent_activities = []
    now = datetime.now()
    
    try:
        # Get recent CVs with optimized query
        recent_cvs = CV.objects.order_by('-processed_at')[:5]
        for cv in recent_cvs:
            recent_activities.append({
                'time': cv.processed_at or now,
                'description': f'New CV uploaded: {cv.name or "N/A"}'
            })
    except Exception as e:
        print(f"Error fetching recent CVs: {e}")
        
    try:
        # Get recent jobs with optimized query
        recent_jobs = JobDescription.objects.order_by('-created_at')[:5]
        for job in recent_jobs:
            recent_activities.append({
                'time': job.created_at or now,
                'description': f'New job added: {job.title or "N/A"}'
            })
    except Exception as e:
        print(f"Error fetching recent jobs: {e}")
        
    try:
        # Get recent matches with optimized query using select_related
        recent_matches = MatchResult.objects.select_related('cv', 'job').order_by('-matched_at')[:5]
        for match in recent_matches:
            cv_name = match.cv.name if match.cv else "N/A"
            job_title = match.job.title if match.job else "N/A"
            recent_activities.append({
                'time': match.matched_at or now,
                'description': f'Matched CV {cv_name} with job {job_title} (Score: {int(match.total_score * 100)}%)'
            })
    except Exception as e:
        print(f"Error fetching recent matches: {e}")
        
    # Sort activities by time
    recent_activities.sort(key=lambda x: x.get('time', datetime.min), reverse=True)
    
    # Limit to 10 most recent
    return recent_activities[:10]